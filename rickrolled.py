 #rolled.pyM-D
#This is for fun not used for harm, only. Use at your own risk.
import importlib
import subprocess
import sys
import time
import webbrowser
import os

# -------------------------------------------------
# Dependency auto-check (installs only if missing)
# -------------------------------------------------
def ensure_package(pkg_name, import_name=None):
    if import_name is None:
        import_name = pkg_name
    try:
        importlib.import_module(import_name)
    except ImportError:
        print(f"📦 Installing missing package: {pkg_name}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg_name])

try:
    ensure_package("python-docx", "docx")
    ensure_package("reportlab")
    ensure_package("openpyxl")

    from docx import Document
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    import openpyxl

except Exception as e:
    print(f"⚠️ Dependency error: {e}")
    print("Continuing without reinstall attempts...")
    Document = None
    SimpleDocTemplate = None
    Paragraph = None
    getSampleStyleSheet = None
    openpyxl = None

# -------------------------------------------------
# Rickroll storm
# -------------------------------------------------
rickroll_url = "https://www.youtube.com/clip/UgkxmchFJAVMvKGkgzW"

def make_prank_files():
    # 1. Windows Shortcut (.url)
    with open("Rickroll.url", "w") as f:
        f.write(f"[InternetShortcut]\nURL={rickroll_url}\n")

    # 2. HTML with a link
    with open("rickroll.html", "w") as f:
        f.write(f"""<!DOCTYPE html>
<html><head><title>Special Link</title></head>
<body><h2><a href="{rickroll_url}">🎁 Claim Your Reward</a></h2></body>
</html>""")

    # 3. HTML auto-redirect
    with open("rickroll_redirect.html", "w") as f:
        f.write(f"""<!DOCTYPE html>
<html>
<head><meta http-equiv="refresh" content="0; url={rickroll_url}" /></head>
<body><p>If not redirected, <a href="{rickroll_url}">click here</a>.</p></body>
</html>""")

    # 4. Word doc
    if Document:
        doc = Document()
        p = doc.add_paragraph("Click here: ")
        p.add_run("🎵 Surprise Link 🎵")
        doc.add_paragraph(rickroll_url)
        doc.save("rickroll.docx")

    # 5. PDF
    if SimpleDocTemplate and Paragraph:
        doc_pdf = SimpleDocTemplate("rickroll.pdf")
        styles = getSampleStyleSheet()
        story = [Paragraph(f'<a href="{rickroll_url}">🔥 Exclusive Offer 🔥</a>', styles["Normal"])]
        doc_pdf.build(story)

    # 6. Excel
    if openpyxl:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Rickroll"
        ws["A1"] = "Click me!"
        ws["A1"].hyperlink = rickroll_url
        ws["A1"].style = "Hyperlink"
        wb.save("rickroll.xlsx")

def launch_rickroll_storm():
    files = [
        "Rickroll.url",
        "rickroll.html",
        "rickroll_redirect.html",
        "rickroll.docx",
        "rickroll.pdf",
        "rickroll.xlsx"
    ]

    print("🚨 Escalating Rickroll Storm activated! (Press Ctrl+C to stop)")
    wave = 1
    try:
        while True:
            print(f"\n🌪️ Wave {wave}: launching {2**(wave-1)} direct Rickrolls + files!")
            # Open prank files once per wave
            for f in files:
                if os.path.exists(f):
                    webbrowser.open(f"file://{os.path.abspath(f)}", new=2)

            # Escalating direct links (1, 2, 4, 8, ...)
            for _ in range(2**(wave-1)):
                webbrowser.open(rickroll_url, new=2)

            wave += 1
            time.sleep(8)  # pause before next wave
    except KeyboardInterrupt:
        print("\n😅 The storm has ended... for now.")

# -------------------------------------------------
# Main with failsafe restart
# -------------------------------------------------
if __name__ == "__main__":
    try:
        make_prank_files()
        launch_rickroll_storm()
    except Exception as e:
        print(f"💥 Unexpected error: {e}")
        print("Restarting Rickroll storm without reinstall...")
        time.sleep(2)
        launch_rickroll_storm()
~ $
